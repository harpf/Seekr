"""Extractor unit tests.

Covers the previously-untested text/markdown extractors against real committed
fixtures, and the legacy-office extractor with subprocess monkeypatched so the
test is hermetic (no antiword/catppt install required).
"""
from __future__ import annotations

import subprocess
from pathlib import Path

import pytest

from document_search.extractors.legacy_office_extractor import LegacyOfficeTextExtractor
from document_search.extractors.md_extractor import MdTextExtractor
from document_search.extractors.txt_extractor import TxtTextExtractor

FIXTURES = Path(__file__).parent / "fixtures"


def test_txt_extractor_reads_committed_fixture():
    result = TxtTextExtractor().extract(FIXTURES / "sample.txt")
    assert result.status == "ok"
    assert len(result.blocks) == 1
    block = result.blocks[0]
    assert block.block_type == "text_file"
    assert "Invoice number 12345" in block.text
    assert block.extractor == "TxtTextExtractor"


def test_md_extractor_reads_committed_fixture():
    result = MdTextExtractor().extract(FIXTURES / "sample.md")
    assert result.status == "ok"
    assert len(result.blocks) == 1
    block = result.blocks[0]
    assert block.block_type == "markdown"
    assert "annualreport" in block.text
    assert "Heading One" in block.text


def test_txt_extractor_empty_file_yields_no_blocks(tmp_path: Path):
    empty = tmp_path / "empty.txt"
    empty.write_text("", encoding="utf-8")
    result = TxtTextExtractor().extract(empty)
    assert result.status == "ok"
    assert result.blocks == []


def test_txt_extractor_tolerates_bad_bytes(tmp_path: Path):
    # extract() uses errors="ignore"; invalid UTF-8 must not raise.
    bad = tmp_path / "bad.txt"
    bad.write_bytes(b"hello \xff\xfe world")
    result = TxtTextExtractor().extract(bad)
    assert result.status == "ok"
    assert "hello" in result.blocks[0].text


def test_legacy_office_extractor_success_with_mocked_subprocess(monkeypatch, tmp_path: Path):
    doc = tmp_path / "legacy.doc"
    doc.write_bytes(b"\xd0\xcf\x11\xe0fake-ole")  # OLE magic, content irrelevant

    def fake_run(cmd, **kwargs):
        assert cmd[0] == "antiword"
        assert cmd[1] == str(doc)
        return subprocess.CompletedProcess(cmd, returncode=0, stdout="Extracted legacy text", stderr="")

    monkeypatch.setattr(subprocess, "run", fake_run)
    result = LegacyOfficeTextExtractor().extract(doc)
    assert result.status == "ok"
    assert result.blocks[0].block_type == "legacy_document"
    assert result.blocks[0].text == "Extracted legacy text"


def test_legacy_office_extractor_ppt_uses_catppt(monkeypatch, tmp_path: Path):
    ppt = tmp_path / "slides.ppt"
    ppt.write_bytes(b"\xd0\xcf\x11\xe0fake-ole")
    seen = {}

    def fake_run(cmd, **kwargs):
        seen["cmd0"] = cmd[0]
        return subprocess.CompletedProcess(cmd, returncode=0, stdout="slide text", stderr="")

    monkeypatch.setattr(subprocess, "run", fake_run)
    result = LegacyOfficeTextExtractor().extract(ppt)
    assert seen["cmd0"] == "catppt"
    assert result.status == "ok"


def test_legacy_office_extractor_nonzero_returncode_is_error(monkeypatch, tmp_path: Path):
    doc = tmp_path / "broken.doc"
    doc.write_bytes(b"\xd0\xcf\x11\xe0")

    def fake_run(cmd, **kwargs):
        return subprocess.CompletedProcess(cmd, returncode=2, stdout="", stderr="antiword: cannot parse")

    monkeypatch.setattr(subprocess, "run", fake_run)
    result = LegacyOfficeTextExtractor().extract(doc)
    assert result.status == "error"
    assert "cannot parse" in (result.error_message or "")


def test_legacy_office_extractor_missing_binary_is_error(monkeypatch, tmp_path: Path):
    doc = tmp_path / "x.doc"
    doc.write_bytes(b"\xd0\xcf\x11\xe0")

    def fake_run(cmd, **kwargs):
        raise FileNotFoundError(cmd[0])

    monkeypatch.setattr(subprocess, "run", fake_run)
    result = LegacyOfficeTextExtractor().extract(doc)
    assert result.status == "error"
    assert "not installed" in (result.error_message or "")


@pytest.mark.skipif(
    pytest.importorskip("docx", reason="python-docx not installed") is None,
    reason="python-docx required",
)
def test_docx_extractor_with_generated_fixture(tmp_path: Path):
    from docx import Document

    from document_search.extractors.docx_extractor import DocxTextExtractor

    path = tmp_path / "generated.docx"
    doc = Document()
    doc.add_heading("Quarterly Title", level=1)
    doc.add_paragraph("Body paragraph mentioning invoice 9876.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Key"
    table.rows[0].cells[1].text = "Value"
    doc.save(str(path))

    result = DocxTextExtractor().extract(path)
    assert result.status == "ok"
    all_text = " ".join(b.text for b in result.blocks)
    assert "Quarterly Title" in all_text
    assert "invoice 9876" in all_text
    block_types = {b.block_type for b in result.blocks}
    assert "table" in block_types
