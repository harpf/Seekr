from __future__ import annotations

import os
from pathlib import Path

from docx import Document

from document_search.extractors.base import TextExtractor
from document_search.models import ContentBlock, ExtractionResult
from document_search.services.ocr_service import ocr_office_embedded_images


class DocxTextExtractor(TextExtractor):
    def extract(self, file_path: Path) -> ExtractionResult:
        try:
            doc = Document(str(file_path))
            blocks: list[ContentBlock] = []
            idx = 1
            for paragraph in doc.paragraphs:
                txt = paragraph.text.strip()
                if txt:
                    btype = "heading" if paragraph.style and "Heading" in paragraph.style.name else "paragraph"
                    blocks.append(ContentBlock(btype, idx, txt, self.__class__.__name__, {}))
                    idx += 1
            for table_i, table in enumerate(doc.tables, start=1):
                rows = ["\t".join(cell.text.strip() for cell in row.cells) for row in table.rows]
                txt = "\n".join(r for r in rows if r.strip())
                if txt:
                    blocks.append(ContentBlock("table", table_i, txt, self.__class__.__name__, {}))
            # headers/footers
            sec_idx = 1
            for section in doc.sections:
                h_text = "\n".join(p.text.strip() for p in section.header.paragraphs if p.text.strip())
                f_text = "\n".join(p.text.strip() for p in section.footer.paragraphs if p.text.strip())
                if h_text:
                    blocks.append(ContentBlock("header", sec_idx, h_text, self.__class__.__name__, {}))
                if f_text:
                    blocks.append(ContentBlock("footer", sec_idx, f_text, self.__class__.__name__, {}))
                sec_idx += 1
            if os.getenv("DOCUMENT_SEARCH_OCR_ENABLED", "false").lower() == "true":
                for i, txt in enumerate(ocr_office_embedded_images(file_path), start=1):
                    blocks.append(ContentBlock("ocr_image", i, txt, self.__class__.__name__, {"source": "embedded_image"}))
            cp = doc.core_properties
            meta = {"title": cp.title, "author": cp.author, "category": cp.category, "created": str(cp.created) if cp.created else None}
            return ExtractionResult(file_path=file_path, status="ok", document_metadata={"docx_properties": meta}, blocks=blocks)
        except Exception as ex:
            return ExtractionResult(file_path=file_path, status="error", error_message=str(ex))
